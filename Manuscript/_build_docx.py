"""
Render AGAPP Markdown files to formatted .docx files, complete with embedded images!

Order of sections in Chapter 3:
  1. Project Design (Agile Model, System Architecture)
  2. Flowcharts (Figures 4 - 38)
  3. Wireframes and Interface Layouts (Figures 39 - 62)
  4. Use Case Diagrams (Figures 63 - 66)
  5. Project Development (Agile Model Iterations)
  6. Project Testing and Evaluation Procedures

Format follows the reference capstone (MANUSCRIPT CHAP 1-3 OSAS):
  * letter paper, 1.5in left margin (binding edge), 1in elsewhere
  * double-spaced, justified body with a 0.5in first-line indent
  * chapter titles centred, bold, all caps
  * section headings bold, flush left, not indented
  * page number in the top-right header
  * tables styled with borders, padded cells, and bold headers
  * embedded flowchart, wireframe, and use case PNG images centered above figure captions
"""
import sys
import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

FONT = "Arial"
SIZE = Pt(12)

FLOWCHART_DIR = r"c:\Users\Lawrence\Documents\PROJECTS\AGAP\Docs\Flowcharts\png\flowcharts"
WIREFRAME_DIR = r"c:\Users\Lawrence\Documents\PROJECTS\AGAP\Docs\Wireframes\Wireframe Agapp"

IMAGE_MAP = {
    # -------------------------------------------------------------
    # FLOWCHARTS (Figures 4 - 38)
    # -------------------------------------------------------------
    "Figure 4": os.path.join(FLOWCHART_DIR, "01-Citizen-App-Citizen App — Entry.drawio.png"),
    "Figure 5": os.path.join(FLOWCHART_DIR, "01-Citizen-App-1 — Guest Browsing and Sign Up.drawio.png"),
    "Figure 6": os.path.join(FLOWCHART_DIR, "01-Citizen-App-2 — Main Interface (Tabs).drawio.png"),
    "Figure 7": os.path.join(FLOWCHART_DIR, "01-Citizen-App-3 — Apply for a Service.drawio.png"),
    "Figure 8": os.path.join(FLOWCHART_DIR, "01-Citizen-App-3A — Tracking and Claiming.drawio.png"),
    "Figure 9": os.path.join(FLOWCHART_DIR, "01-Citizen-App-4 — Report an Issue.drawio.png"),
    "Figure 10": os.path.join(FLOWCHART_DIR, "01-Citizen-App-4A — Photo Check and Follow-up.drawio.png"),
    "Figure 11": os.path.join(FLOWCHART_DIR, "01-Citizen-App-5 — Community Forum.drawio.png"),
    "Figure 12": os.path.join(FLOWCHART_DIR, "01-Citizen-App-6 — Profile and Account.drawio.png"),
    "Figure 13": os.path.join(FLOWCHART_DIR, "01-Citizen-App-6A — Deleting the Account.drawio.png"),
    "Figure 14": os.path.join(FLOWCHART_DIR, "01-Citizen-App-7 — News and Announcements.drawio.png"),
    "Figure 15": os.path.join(FLOWCHART_DIR, "01-Citizen-App-8 — Assistant (Chatbot).drawio.png"),
    "Figure 16": os.path.join(FLOWCHART_DIR, "01-Citizen-App-9 — Restricted or Banned Account.drawio.png"),
    "Figure 17": os.path.join(FLOWCHART_DIR, "01-Citizen-App-10 — Identity Verification.drawio.png"),
    
    # LGU Admin Flowcharts
    "Figure 18": os.path.join(FLOWCHART_DIR, "02-LGU-Admin-LGU Admin — Main Flow.drawio.png"),
    "Figure 19": os.path.join(FLOWCHART_DIR, "02-LGU-Admin-1 — Dashboard.drawio.png"),
    "Figure 20": os.path.join(FLOWCHART_DIR, "02-LGU-Admin-2 — Issue Reports.drawio.png"),
    "Figure 21": os.path.join(FLOWCHART_DIR, "02-LGU-Admin-3 — Service Requests.drawio.png"),
    "Figure 22": os.path.join(FLOWCHART_DIR, "02-LGU-Admin-4 — eServices Catalog.drawio.png"),
    "Figure 23": os.path.join(FLOWCHART_DIR, "02-LGU-Admin-5 — Community and News.drawio.png"),
    "Figure 24": os.path.join(FLOWCHART_DIR, "02-LGU-Admin-6 — Forum Moderation.drawio.png"),
    "Figure 25": os.path.join(FLOWCHART_DIR, "02-LGU-Admin-7 — Facilities.drawio.png"),
    "Figure 26": os.path.join(FLOWCHART_DIR, "02-LGU-Admin-8 — Citizen Guide.drawio.png"),
    "Figure 27": os.path.join(FLOWCHART_DIR, "02-LGU-Admin-9 — Citizens and Moderation.drawio.png"),
    "Figure 28": os.path.join(FLOWCHART_DIR, "02-LGU-Admin-9A — Reviewing a Citizen Appeal.drawio.png"),
    "Figure 29": os.path.join(FLOWCHART_DIR, "02-LGU-Admin-10 — ID Verifications.drawio.png"),
    "Figure 30": os.path.join(FLOWCHART_DIR, "02-LGU-Admin-11 — Settings and Staff.drawio.png"),
    "Figure 31": os.path.join(FLOWCHART_DIR, "02-LGU-Admin-12 — Logout.drawio.png"),

    # Super Admin Flowcharts
    "Figure 32": os.path.join(FLOWCHART_DIR, "03-Super-Admin-Super Admin — Main Flow.drawio.png"),
    "Figure 33": os.path.join(FLOWCHART_DIR, "03-Super-Admin-1 — Cross-LGU Dashboard.drawio.png"),
    "Figure 34": os.path.join(FLOWCHART_DIR, "03-Super-Admin-2 — LGU Directory.drawio.png"),
    "Figure 35": os.path.join(FLOWCHART_DIR, "03-Super-Admin-2A — Onboarding a New LGU (Wizard).drawio.png"),
    "Figure 36": os.path.join(FLOWCHART_DIR, "03-Super-Admin-3 — Analytics.drawio.png"),
    "Figure 37": os.path.join(FLOWCHART_DIR, "03-Super-Admin-4 — Super Admin Settings.drawio.png"),
    "Figure 38": os.path.join(FLOWCHART_DIR, "03-Super-Admin-5 — Logout.drawio.png"),

    # -------------------------------------------------------------
    # WIREFRAMES (Figures 39 - 62)
    # -------------------------------------------------------------
    "Figure 39": os.path.join(WIREFRAME_DIR, "Select your lgu Mobile.png"),
    "Figure 40": os.path.join(WIREFRAME_DIR, "Login.Welcome Mobile.png"),
    "Figure 41": os.path.join(WIREFRAME_DIR, "Homepage Mobile.png"),
    "Figure 42": os.path.join(WIREFRAME_DIR, "Services Mobile.png"),
    "Figure 43": os.path.join(WIREFRAME_DIR, "Issue Report.png"),
    "Figure 44": os.path.join(WIREFRAME_DIR, "Reports Mobile.png"),
    "Figure 45": os.path.join(WIREFRAME_DIR, "Map Mobile.png"),
    "Figure 46": os.path.join(WIREFRAME_DIR, "Community Mobile.png"),
    "Figure 47": os.path.join(WIREFRAME_DIR, "Forum Mobile.png"),
    "Figure 48": os.path.join(WIREFRAME_DIR, "Chatbot Mobile.png"),
    "Figure 49": os.path.join(WIREFRAME_DIR, "Profile Mobile.png"),
    "Figure 50": os.path.join(WIREFRAME_DIR, "Citizen Verification.png"),
    "Figure 51": os.path.join(WIREFRAME_DIR, "Dashboard Overview Agapp Portal.png"),
    "Figure 52": os.path.join(WIREFRAME_DIR, "Service Request Liliw laguna.png"),
    "Figure 53": os.path.join(WIREFRAME_DIR, "Body.png"),
    "Figure 54": os.path.join(WIREFRAME_DIR, "Citizen and Moderation.png"),
    "Figure 55": os.path.join(WIREFRAME_DIR, "Citizen Guide Directory.png"),
    "Figure 56": os.path.join(WIREFRAME_DIR, "Community Liliw Laguna.png"),
    "Figure 57": os.path.join(WIREFRAME_DIR, "Forum Moderation.png"),
    "Figure 58": os.path.join(WIREFRAME_DIR, "Facilities Map.png"),
    "Figure 59": os.path.join(WIREFRAME_DIR, "eServices Catalog.png"),
    "Figure 60": os.path.join(WIREFRAME_DIR, "Settings Liliw Laguna.png"),
    "Figure 61": os.path.join(WIREFRAME_DIR, "LGU Directory.png"),
    "Figure 62": os.path.join(WIREFRAME_DIR, "System settings.png"),

    # -------------------------------------------------------------
    # USE CASE DIAGRAMS (Figures 63 - 66)
    # -------------------------------------------------------------
    "Figure 63": os.path.join(FLOWCHART_DIR, "04-Use-Case-Diagrams-Figure — Use Case Diagram_ Super Administrator.drawio.png"),
    "Figure 64": os.path.join(FLOWCHART_DIR, "04-Use-Case-Diagrams-Figure — Use Case Diagram_ LGU Administrator.drawio.png"),
    "Figure 65": os.path.join(FLOWCHART_DIR, "04-Use-Case-Diagrams-Figure — Use Case Diagram_ LGU Personnel.drawio.png"),
    "Figure 66": os.path.join(FLOWCHART_DIR, "04-Use-Case-Diagrams-Figure — Use Case Diagram_ Verified Citizen.drawio.png"),
}

def add_page_number_header(section):
    """Right-aligned live page-number field in the header."""
    p = section.header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    for instr, kind in (("begin", "fldCharType"), (None, "instrText"), ("end", "fldCharType")):
        if kind == "instrText":
            el = OxmlElement("w:instrText")
            el.set(qn("xml:space"), "preserve")
            el.text = "PAGE"
        else:
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), instr)
        run._r.append(el)
    run.font.name = FONT
    run.font.size = SIZE

def style_doc(doc):
    st = doc.styles["Normal"]
    st.font.name = FONT
    st.font.size = SIZE
    st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    pf = st.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)

    for s in doc.sections:
        s.left_margin = Inches(1.5)
        s.right_margin = Inches(1.0)
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        add_page_number_header(s)

BOLD_RE = re.compile(r"\*\*(.+?)\*\*")

def emit_runs(par, text):
    """Write text into a paragraph, honouring **bold** spans."""
    pos = 0
    clean_text = text.replace("<br>", "\n").replace("<br/>", "\n")
    for m in BOLD_RE.finditer(clean_text):
        if m.start() > pos:
            par.add_run(clean_text[pos:m.start()])
        par.add_run(m.group(1)).bold = True
        pos = m.end()
    if pos < len(clean_text):
        par.add_run(clean_text[pos:])

def para(doc, text, *, align=None, indent=None, left=None, bold=False, caps=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    if indent is not None:
        pf.first_line_indent = indent
    if left is not None:
        pf.left_indent = left
    if bold:
        r = p.add_run(text.upper() if caps else text)
        r.bold = True
    else:
        emit_runs(p, text)
    return p

NUM_RE = re.compile(r"^(\d+)\.\s+(.*)$")
LET_RE = re.compile(r"^([a-z])\)\s+(.*)$")

def set_cell_border(cell, **kwargs):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            element.set(qn('w:val'), edge_data.get('val', 'single'))
            element.set(qn('w:sz'), str(edge_data.get('sz', 4)))
            element.set(qn('w:space'), str(edge_data.get('space', 0)))
            element.set(qn('w:color'), edge_data.get('color', 'auto'))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def render_table(doc, table_lines):
    if not table_lines:
        return
    rows_data = []
    for line in table_lines:
        parts = [c.strip() for c in line.strip('|').split('|')]
        if all(re.match(r'^:?-+:?$', p) for p in parts if p):
            continue
        rows_data.append(parts)

    if not rows_data:
        return

    num_rows = len(rows_data)
    num_cols = max(len(r) for r in rows_data)

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    border_spec = dict(sz=4, val='single', color='CCCCCC')

    for r_idx, row_cells in enumerate(rows_data):
        row = table.rows[r_idx]
        for c_idx, cell_value in enumerate(row_cells):
            if c_idx < num_cols:
                cell = row.cells[c_idx]
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                set_cell_border(cell, top=border_spec, bottom=border_spec, left=border_spec, right=border_spec)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)
                if r_idx == 0:
                    emit_runs(p, f"**{cell_value}**")
                else:
                    emit_runs(p, cell_value)

    doc.add_paragraph()

FIG_MATCH_RE = re.compile(r"Figure\s+(\d+)", re.IGNORECASE)

def try_insert_image_for_figure(doc, figure_text):
    m = FIG_MATCH_RE.search(figure_text)
    if not m:
        return False
    fig_key = f"Figure {m.group(1)}"
    img_path = IMAGE_MAP.get(fig_key)
    if img_path and os.path.exists(img_path):
        img_p = doc.add_paragraph()
        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_p.paragraph_format.space_before = Pt(12)
        img_p.paragraph_format.space_after = Pt(6)
        run = img_p.add_run()
        max_w = Inches(3.0) if ("Mobile" in os.path.basename(img_path) or "Verification.png" in os.path.basename(img_path)) else Inches(5.2)
        run.add_picture(img_path, width=max_w)
        return True
    return False

def build(src_file, out_file):
    lines = open(src_file, encoding="utf-8").read().split("\n")
    doc = Document()
    style_doc(doc)

    i = 0
    n = len(lines)
    while i < n:
        raw = lines[i]
        line = raw.rstrip()
        s = line.strip()

        if not s:
            i += 1
            continue

        if s == "---":
            doc.add_page_break()
            i += 1
            continue

        if s.startswith("|") and s.endswith("|"):
            table_lines = []
            while i < n and lines[i].strip().startswith("|") and lines[i].strip().endswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            render_table(doc, table_lines)
            continue

        if s.startswith("#### "):
            para(doc, s[5:], bold=True)
            i += 1
            continue
        if s.startswith("### "):
            para(doc, s[4:], bold=True)
            i += 1
            continue
        if s.startswith("## "):
            para(doc, s[3:], align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, caps=True)
            i += 1
            continue
        if s.startswith("# "):
            para(doc, s[2:], align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, caps=True)
            i += 1
            continue

        if s.startswith("**[Figure") or s.startswith("[Figure") or s.startswith("**Figure"):
            fig_clean = s.strip("*[]")
            inserted = try_insert_image_for_figure(doc, fig_clean)
            para(doc, f"[{fig_clean}]" if not fig_clean.startswith("Figure") else fig_clean, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
            i += 1
            continue

        m = NUM_RE.match(s)
        if m:
            para(doc, f"{m.group(1)}. {m.group(2)}",
                 align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                 left=Inches(0.5), indent=Inches(-0.25))
            i += 1
            continue

        m = LET_RE.match(s)
        if m:
            para(doc, f"{m.group(1)}) {m.group(2)}",
                 align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                 left=Inches(1.0), indent=Inches(-0.25))
            i += 1
            continue

        # ordinary body paragraph
        para(doc, s, align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=Inches(0.5))
        i += 1

    doc.save(out_file)
    print(f"Successfully generated with embedded images: {out_file}")

if __name__ == "__main__":
    src = sys.argv[1] if len(sys.argv) > 1 else "AGAPP-Chapter-1-2-REVISED.md"
    out = sys.argv[2] if len(sys.argv) > 2 else "AGAPP-Chapter-1-2-REVISED.docx"
    build(src, out)
